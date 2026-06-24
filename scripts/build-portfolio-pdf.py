"""
Assemble the entire portfolio (about + all work entries) into a single
Markdown document, then convert to DOCX and PDF for offline distribution.

Output: public/Adesola-Gabriel-Adeola-Portfolio.pdf (and .docx)

Requires:
    pip install python-docx docx2pdf gray-matter-py

On Windows/macOS, docx2pdf uses Word/Pages under the hood.
"""

from __future__ import annotations

import argparse
import platform
import re
import sys
from pathlib import Path


REPO_ROOT = Path(__file__).resolve().parent.parent
ABOUT_MDX = REPO_ROOT / "content" / "about.mdx"
WORK_DIR = REPO_ROOT / "content" / "work"
PUBLIC_DIR = REPO_ROOT / "public"

OUTPUT_STEM = "Adesola-Gabriel-Adeola-Portfolio"

HEADER = """# Adesola Gabriel Adeola

Security Engineer · Columbus, Ohio
aadeola20@outlook.com · linkedin.com/in/adesola-adeola-5255841b7 · github.com/Godson90

This document is the offline edition of my portfolio. It mirrors the website at gabrieladeola.dev: an about page, a how-I-work statement, and write-ups of every project listed under selected work.

---
"""


def strip_frontmatter(text: str) -> tuple[dict[str, str], str]:
    """Split YAML frontmatter from body. Returns (front, body)."""
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    front_block = text[3:end].strip()
    body = text[end + 4:].lstrip("\n")
    front: dict[str, str] = {}
    for line in front_block.splitlines():
        if ":" in line:
            key, _, value = line.partition(":")
            front[key.strip()] = value.strip().strip('"').strip("'")
    return front, body


def load_about() -> str:
    raw = ABOUT_MDX.read_text(encoding="utf-8")
    _, body = strip_frontmatter(raw)
    return body.strip()


def load_work_entries() -> list[tuple[dict[str, str], str]]:
    entries: list[tuple[dict[str, str], str]] = []
    for mdx_path in sorted(WORK_DIR.glob("*.mdx")):
        front, body = strip_frontmatter(mdx_path.read_text(encoding="utf-8"))
        entries.append((front, body.strip()))
    entries.sort(key=lambda pair: pair[0].get("catalogId", "99"))
    return entries


def assemble_markdown() -> str:
    chunks: list[str] = [HEADER.strip(), ""]

    chunks.append("# About")
    chunks.append("")
    chunks.append(load_about())
    chunks.append("")
    chunks.append("---")
    chunks.append("")

    chunks.append("# Selected Work")
    chunks.append("")
    for front, body in load_work_entries():
        title = front.get("title", front.get("slug", "Untitled"))
        one_liner = front.get("oneLiner", "")
        year = front.get("year", "")
        role = front.get("role", "")
        catalog_id = front.get("catalogId", "")
        header_line = " · ".join(filter(None, [f"#{catalog_id}" if catalog_id else "", year, role]))
        chunks.append(f"## {title}")
        if header_line:
            chunks.append(f"*{header_line}*")
        if one_liner:
            chunks.append(f"_{one_liner}_")
        chunks.append("")
        chunks.append(_normalize_body(body))
        chunks.append("")
        chunks.append("---")
        chunks.append("")
    return "\n".join(chunks)


def _normalize_body(body: str) -> str:
    """Demote one heading level so the merged doc stays hierarchical.

    Each work MDX uses `##` for its sub-sections; promote those to `###` so
    they sit below the project's `##` title in the merged document.
    """
    out_lines: list[str] = []
    for line in body.splitlines():
        if line.startswith("######"):
            out_lines.append(line)
        elif line.startswith("##"):
            out_lines.append("#" + line)
        else:
            out_lines.append(line)
    return "\n".join(out_lines)


def parse_markdown(md_text: str) -> list[tuple[str, str]]:
    tokens: list[tuple[str, str]] = []
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            tokens.append(("blank", ""))
            continue
        if line.strip() == "---":
            tokens.append(("hr", ""))
            continue
        if line.startswith("#### "):
            tokens.append(("h4", _strip_inline(line[5:])))
        elif line.startswith("### "):
            tokens.append(("h3", _strip_inline(line[4:])))
        elif line.startswith("## "):
            tokens.append(("h2", _strip_inline(line[3:])))
        elif line.startswith("# "):
            tokens.append(("h1", _strip_inline(line[2:])))
        elif re.match(r"^[-*+]\s+", line):
            tokens.append(("bullet", _strip_inline(re.sub(r"^[-*+]\s+", "", line))))
        else:
            tokens.append(("para", _strip_inline(line)))
    return tokens


def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def build_docx(tokens: list[tuple[str, str]], out_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        sys.exit("python-docx is not installed. Run: pip install python-docx docx2pdf")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, text in tokens:
        if kind == "h1":
            doc.add_heading(text, level=0)
        elif kind == "h2":
            doc.add_heading(text, level=1)
        elif kind == "h3":
            doc.add_heading(text, level=2)
        elif kind == "h4":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "para":
            doc.add_paragraph(text)
        elif kind == "hr":
            doc.add_paragraph("―" * 30)
        elif kind == "blank":
            continue

    doc.save(str(out_path))


def build_pdf(docx_path: Path, pdf_path: Path) -> bool:
    system = platform.system()
    if system not in ("Windows", "Darwin"):
        print(f"PDF conversion not supported on {system}.", file=sys.stderr)
        return False
    try:
        from docx2pdf import convert
    except ImportError:
        print("docx2pdf is not installed. Run: pip install docx2pdf", file=sys.stderr)
        return False
    convert(str(docx_path), str(pdf_path))
    return True


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a single PDF of the entire portfolio.")
    parser.add_argument("--no-pdf", action="store_true", help="Skip PDF generation.")
    parser.add_argument("--out-dir", default=str(PUBLIC_DIR), help="Output directory.")
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    markdown = assemble_markdown()
    md_path = out_dir / f"{OUTPUT_STEM}.md"
    md_path.write_text(markdown, encoding="utf-8")
    print(f"Wrote {md_path} ({len(markdown):,} chars)")

    tokens = parse_markdown(markdown)
    docx_path = out_dir / f"{OUTPUT_STEM}.docx"
    build_docx(tokens, docx_path)
    print(f"Wrote {docx_path}")

    if not args.no_pdf:
        pdf_path = out_dir / f"{OUTPUT_STEM}.pdf"
        if build_pdf(docx_path, pdf_path):
            print(f"Wrote {pdf_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
